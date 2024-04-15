from pypdf import PdfReader
from fastapi_mail import FastMail, ConnectionConfig, MessageSchema
import os
import docx

conf = ConnectionConfig(
    MAIL_USERNAME=os.getenv('MAIL_USERNAME'),
    MAIL_PASSWORD=os.getenv('MAIL_PASSWORD'),
    MAIL_FROM=os.getenv('MAIL_FROM'),
    MAIL_PORT=os.getenv('MAIL_PORT'),
    MAIL_SERVER=os.getenv('MAIL_SERVER'),
    MAIL_FROM_NAME=os.getenv('MAIL_FROM_NAME'),
    USE_CREDENTIALS=True,
    TEMPLATE_FOLDER='./app/templates',
    MAIL_STARTTLS=False,
    MAIL_SSL_TLS=False
)

async def extract_text_from_pdf(file):
    extract_text = ""
    reader = PdfReader(file)
    for i in reader.pages:
        extract_text += i.extract_text()+" "
    return extract_text

async def extract_text_from_docx(file):
    doc = docx.Document(file._file)
    document_text = ""
    for i in doc.paragraphs:
        document_text += i.text

    return document_text

async def extract_text_from_doc(file):
    doc = docx.Document(file._file)
    document_text = ""
    for i in doc.paragraphs:
        document_text += i.text

    return document_text

async def send_email(subject: str, email_to: str, body: dict):
    message = MessageSchema(
        subject=subject,
        recipients=[email_to],
        body=body,
        subtype='html',
    )
    
    fm = FastMail(conf)
    await fm.send_message(message, template_name='templates/send_email.html')
